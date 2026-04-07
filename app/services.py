from __future__ import annotations

import json
import re
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from sqlalchemy.orm import Session

from .models import Artifact, AuditEvent, Document, DocumentVersion, Matter, User

DATE_PATTERNS = [
    re.compile(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b"),
    re.compile(r"\b\d{1,2}\s+[A-Za-z]{3,9}\s+\d{4}\b"),
    re.compile(r"\b[A-Za-z]{3,9}\s+\d{1,2},\s+\d{4}\b"),
]


def log_event(db: Session, user: User, action: str, entity_type: str, entity_id: str, meta: dict | None = None) -> None:
    event = AuditEvent(
        tenant=user.tenant,
        user_id=user.id,
        action=action,
        entity_type=entity_type,
        entity_id=entity_id,
        meta=json.dumps(meta or {}),
    )
    db.add(event)
    db.commit()


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        if suffix == ".docx":
            doc = DocxDocument(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        if suffix in {".txt", ".md"}:
            return path.read_text(errors="ignore")
    except Exception:
        return ""
    return ""


def latest_versions_by_matter(db: Session, matter_id: int) -> list[DocumentVersion]:
    docs = db.query(Document).filter(Document.matter_id == matter_id).all()
    versions: list[DocumentVersion] = []
    for doc in docs:
        version = (
            db.query(DocumentVersion)
            .filter(DocumentVersion.document_id == doc.id)
            .order_by(DocumentVersion.version_number.desc())
            .first()
        )
        if version:
            versions.append(version)
    return versions


def source_refs(db: Session, versions: list[DocumentVersion], max_items: int = 8) -> list[dict]:
    refs = []
    for v in versions[:max_items]:
        doc = db.query(Document).filter(Document.id == v.document_id).first()
        if not doc:
            continue
        refs.append({
            "document_id": doc.id,
            "document_title": doc.title,
            "tag": doc.tag,
            "version": v.version_number,
        })
    return refs


def build_brief(matter: Matter, versions: list[DocumentVersion], db: Session) -> tuple[str, list[dict]]:
    chunks = [v.extracted_text for v in versions if v.extracted_text.strip()]
    combined = "\n".join(chunks)
    sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", combined) if s.strip()]
    top_sentences = sentences[:12]
    content = (
        f"Matter: {matter.title}\n"
        f"Forum: {matter.forum}\n\n"
        "Facts Summary:\n"
        + "\n".join(f"- {s}" for s in top_sentences[:6])
        + "\n\nParty Positions:\n"
        + "- Revenue/Authority position extracted from notice and orders.\n"
        + "- Assessee/Respondent position to be reviewed and refined.\n\n"
        + "Disputed Points:\n"
        + "\n".join(f"- {s}" for s in top_sentences[6:10])
        + "\n\nRelief Sought:\n- Set aside adverse findings and grant appropriate relief based on records."
    )
    return content, source_refs(db, versions)


def build_chronology(versions: list[DocumentVersion], db: Session) -> tuple[str, list[dict]]:
    events = []
    for v in versions:
        text = v.extracted_text or ""
        for pattern in DATE_PATTERNS:
            for m in pattern.finditer(text):
                snippet = text[max(0, m.start() - 80): m.end() + 120].replace("\n", " ")
                doc = db.query(Document).filter(Document.id == v.document_id).first()
                events.append((m.group(0), snippet.strip(), doc.title if doc else f"Document {v.document_id}"))
    unique = []
    seen = set()
    for d, s, t in events:
        key = (d, s[:80], t)
        if key not in seen:
            seen.add(key)
            unique.append((d, s, t))
    unique = unique[:30]
    body = "\n".join(f"- {d}: {s} (Source: {t})" for d, s, t in unique) or "- No explicit date events found."
    return f"Chronology\n{body}", source_refs(db, versions)


def build_issues(versions: list[DocumentVersion], db: Session) -> tuple[str, list[dict]]:
    keywords = ["disallow", "penalty", "mismatch", "suppression", "violation", "unexplained", "non-compliance", "denied"]
    rows = []
    for v in versions:
        doc = db.query(Document).filter(Document.id == v.document_id).first()
        text = (v.extracted_text or "").lower()
        for kw in keywords:
            idx = text.find(kw)
            if idx >= 0:
                snippet = (v.extracted_text or "")[max(0, idx - 80): idx + 150].replace("\n", " ")
                rows.append((kw, snippet.strip(), doc.title if doc else f"Document {v.document_id}"))
    if not rows:
        content = "Issue List\n- No keyword-based issues auto-detected. Add manual issues during review."
    else:
        content = "Issue List\n" + "\n".join(
            f"- Issue: {kw.title()}\n  Evidence: {snip}\n  Source: {src}" for kw, snip, src in rows[:20]
        )
    return content, source_refs(db, versions)


def build_draft_response(matter: Matter, versions: list[DocumentVersion], db: Session) -> tuple[str, list[dict]]:
    issue_content, _ = build_issues(versions, db)
    lines = issue_content.splitlines()
    issue_lines = [l.replace("- Issue: ", "") for l in lines if l.startswith("- Issue:")][:8]
    if not issue_lines:
        issue_lines = ["Facts and legal interpretation as per records"]
    draft = (
        f"Draft Reply / Written Submissions\n"
        f"In the matter of: {matter.title}\nForum: {matter.forum}\n\n"
        "1. Preliminary submissions\n"
        "The present reply is being filed based on available records and annexed evidence.\n\n"
        "2. Brief facts\n"
        "The facts are set out in the accompanying chronology and document brief.\n\n"
        "3. Issues for adjudication\n"
        + "\n".join(f"- {i}" for i in issue_lines)
        + "\n\n4. Prayer\nIt is respectfully prayed that the proceedings/adverse conclusions be dropped and relief be granted."
    )
    return draft, source_refs(db, versions)


def build_annexure_index(versions: list[DocumentVersion], db: Session) -> tuple[str, list[dict]]:
    items = []
    for i, v in enumerate(versions, start=1):
        doc = db.query(Document).filter(Document.id == v.document_id).first()
        if doc:
            items.append(f"Annexure-{chr(64 + min(i, 26))}: {doc.title} (tag: {doc.tag}, v{v.version_number})")
    text = "Annexure Index\n" + ("\n".join(f"- {i}" for i in items) if items else "- No annexures uploaded.")
    return text, source_refs(db, versions)


def create_artifact(db: Session, matter_id: int, artifact_type: str, title: str, content: str, sources: list[dict], author_id: int) -> Artifact:
    last = (
        db.query(Artifact)
        .filter(Artifact.matter_id == matter_id, Artifact.artifact_type == artifact_type)
        .order_by(Artifact.version_number.desc())
        .first()
    )
    version = 1 if not last else last.version_number + 1
    artifact = Artifact(
        matter_id=matter_id,
        artifact_type=artifact_type,
        version_number=version,
        title=title,
        content=content,
        sources_json=json.dumps(sources),
        author_id=author_id,
    )
    db.add(artifact)
    db.commit()
    db.refresh(artifact)
    return artifact


def write_docx(content: str, path: Path) -> None:
    doc = DocxDocument()
    for line in content.splitlines():
        doc.add_paragraph(line)
    doc.save(path)


def _build_text_pdf(path: Path, title: str, body: str) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, title)
    y -= 30
    c.setFont("Helvetica", 10)
    for raw in body.splitlines():
        line = raw[:130]
        c.drawString(50, y, line)
        y -= 14
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - 50
    c.save()


def build_filing_bundle_pdf(
    output_path: Path,
    matter: Matter,
    artifacts: list[Artifact],
    latest_versions: list[DocumentVersion],
    db: Session,
) -> None:
    temp_dir = output_path.parent
    writer = PdfWriter()

    cover_path = temp_dir / f"cover_{datetime.utcnow().timestamp()}.pdf"
    _build_text_pdf(cover_path, f"Filing Bundle - {matter.title}", f"Forum: {matter.forum}\nGenerated: {datetime.utcnow().isoformat()}")
    for p in PdfReader(str(cover_path)).pages:
        writer.add_page(p)

    for a in artifacts:
        ap = temp_dir / f"artifact_{a.id}_{datetime.utcnow().timestamp()}.pdf"
        _build_text_pdf(ap, f"{a.title} (v{a.version_number})", a.content)
        for p in PdfReader(str(ap)).pages:
            writer.add_page(p)

    for idx, v in enumerate(latest_versions, start=1):
        doc = db.query(Document).filter(Document.id == v.document_id).first()
        label = doc.title if doc else f"Document {v.document_id}"
        fp = Path(v.file_path)
        if fp.suffix.lower() == ".pdf" and fp.exists():
            try:
                for p in PdfReader(str(fp)).pages:
                    writer.add_page(p)
                continue
            except Exception:
                pass
        placeholder = temp_dir / f"annex_{idx}_{datetime.utcnow().timestamp()}.pdf"
        _build_text_pdf(placeholder, f"Annexure-{chr(64 + min(idx, 26))}", f"{label}\nOriginal file: {v.file_path}")
        for p in PdfReader(str(placeholder)).pages:
            writer.add_page(p)

    with output_path.open("wb") as f:
        writer.write(f)


def quick_keywords(text: str, topn: int = 12) -> list[str]:
    terms = re.findall(r"\b[a-zA-Z]{4,}\b", text.lower())
    stop = {"that", "this", "with", "from", "have", "shall", "where", "there", "under", "which", "notice", "order"}
    terms = [t for t in terms if t not in stop]
    return [w for w, _ in Counter(terms).most_common(topn)]
